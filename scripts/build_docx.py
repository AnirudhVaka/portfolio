#!/usr/bin/env python3
"""
Python fallback DOCX generator.

Equivalent to scripts/build-docx.ts but uses the python-docx library
(pre-installed in many Linux distros, including this sandbox) so we
can produce the 10 region-specific .docx files WITHOUT running npm
install. Useful when the JS toolchain isn't available, or for CI
images that have Python but not Node's `docx` package.

Reads the typed TS data via scripts/dump-data.mjs (which uses Node 22's
--experimental-strip-types), so the source of truth remains
data/resume.ts and lib/resumeRules.ts — no parallel data file.

Generated artifacts:
   public/downloads/anirudh-vaka-resume-{slug}.docx  (one per region)

Run:
   node --experimental-strip-types scripts/dump-data.mjs > /tmp/resume-data.json
   python3 scripts/build_docx.py /tmp/resume-data.json public/downloads/
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Visual constants ──────────────────────────────────────────────────
ACCENT = RGBColor(0x1A, 0x56, 0xDB)
TEXT = RGBColor(0x1E, 0x1E, 0x2E)
MUTED = RGBColor(0x4A, 0x4A, 0x6A)
RULE_GREY = "D0D5E8"
FONT_FAMILY = "Calibri"


def set_run(run, *, size_pt=11, bold=False, italic=False, color=MUTED, font=FONT_FAMILY):
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_para(doc, *, space_before=0, space_after=2, line_spacing=1.15):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    return p


def add_hyperlink(paragraph, url, text, *, color=ACCENT, size_pt=10):
    """python-docx has no native hyperlink helper — build the relationship by hand."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Font + size.
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_FAMILY)
    rFonts.set(qn("w:hAnsi"), FONT_FAMILY)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_pt * 2))
    rPr.append(sz)

    # Colour.
    col = OxmlElement("w:color")
    col.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(col)

    # Underline.
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def add_bottom_border(paragraph, color=RULE_GREY, size=6):
    """Adds a single-line bottom border to the paragraph — used on section titles."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(doc, text):
    p = add_para(doc, space_before=8, space_after=2)
    run = p.add_run(text.upper())
    set_run(run, size_pt=10, bold=True, color=ACCENT)
    add_bottom_border(p)
    return p


def name_heading(doc, name):
    p = add_para(doc, space_after=2)
    run = p.add_run(name)
    set_run(run, size_pt=22, bold=True, color=TEXT)
    return p


def title_chip(doc, title):
    p = add_para(doc, space_after=4)
    run = p.add_run(title)
    set_run(run, size_pt=11, bold=True, color=ACCENT)
    return p


def contact_line(doc, parts):
    """`parts` is a list of (text, href|None) tuples."""
    p = add_para(doc, space_after=2)
    for i, (text, href) in enumerate(parts):
        if href:
            add_hyperlink(p, href, text, size_pt=10)
        else:
            run = p.add_run(text)
            set_run(run, size_pt=10, color=MUTED)
        if i < len(parts) - 1:
            sep = p.add_run("  |  ")
            set_run(sep, size_pt=10, color=MUTED)
    return p


def meta_line(doc, text):
    p = add_para(doc, space_after=2)
    run = p.add_run(text)
    set_run(run, size_pt=10, italic=True, color=MUTED)
    return p


def body(doc, text):
    p = add_para(doc, space_after=2)
    run = p.add_run(text)
    set_run(run, size_pt=11, color=MUTED)
    return p


def job_header(doc, title, dates):
    """Two-column layout: title left-aligned, dates right-aligned via tab stop."""
    p = add_para(doc, space_before=4, space_after=0)
    pf = p.paragraph_format
    # 6.5 inches = right edge for letter; we use 7" for A4-ish.
    pf.tab_stops.add_tab_stop(Inches(6.6), alignment=WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(title)
    set_run(r1, size_pt=12, bold=True, color=TEXT)
    r2 = p.add_run("\t" + dates)
    set_run(r2, size_pt=11, color=MUTED)
    return p


def company_line(doc, text):
    p = add_para(doc, space_after=2)
    run = p.add_run(text)
    set_run(run, size_pt=11, bold=True, color=ACCENT)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(1)
    pf.line_spacing = 1.2
    pf.left_indent = Inches(0.25)
    run = p.runs[0] if p.runs else p.add_run("")
    if not p.runs:
        run = p.add_run(text)
    else:
        # python-docx populates the first run on style="List Bullet"; replace text.
        run.text = text
    set_run(run, size_pt=11, color=MUTED)
    return p


def project_header(doc, name, url, role, dates):
    p = add_para(doc, space_before=4, space_after=0)
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(6.6), alignment=WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(f"{role} — ")
    set_run(r1, size_pt=12, bold=True, color=TEXT)
    add_hyperlink(p, url, name, size_pt=12)
    host = url.replace("https://", "").replace("http://", "")
    r2 = p.add_run(f"  ({host})")
    set_run(r2, size_pt=10, color=MUTED)
    r3 = p.add_run("\t" + dates)
    set_run(r3, size_pt=11, color=MUTED)
    return p


# ── Section renderers ─────────────────────────────────────────────────


def bullet_allowed(priority, bullet_filter):
    return bullet_filter == "all" or priority == "core"


def render_personal_details(doc, data, rules):
    section_heading(doc, "Personal Details")
    rows = [("Location", rules["locationLine"])]
    if rules["includeNationality"] and rules.get("nationality"):
        rows.append(("Nationality", rules["nationality"]))
    if rules["includeNoticePeriod"]:
        rows.append(("Notice period", data["noticePeriod"]))
    rows.append(("Languages", ", ".join(l["name"] for l in data["languages"])))

    for label, value in rows:
        p = add_para(doc, space_after=2)
        r1 = p.add_run(f"{label}: ")
        set_run(r1, size_pt=11, bold=True, color=TEXT)
        r2 = p.add_run(value)
        set_run(r2, size_pt=11, color=MUTED)


def render_summary(doc, data, rules):
    section_heading(doc, rules["summaryLabel"])
    body(doc, data["summary"])


def render_experience(doc, data, rules):
    section_heading(doc, rules["experienceLabel"])
    for job in data["experience"]:
        bullets = [
            b for b in job["bullets"] if bullet_allowed(b["priority"], rules["bulletFilter"])
        ]
        if not bullets:
            continue
        job_header(doc, job["title"], job["dates"])
        suffix = ""
        if job.get("clientContext"):
            suffix = f" · {job['clientContext']}"
        loc_mode = (
            job["mode"]
            if job["location"].lower() == job["mode"].lower()
            else f"{job['location']} · {job['mode']}"
        )
        company_line(doc, f"{job['company']}{suffix} · {loc_mode}")
        for b in bullets:
            bullet(doc, b["text"])


def render_side_projects(doc, data, rules):
    section_heading(doc, rules["projectsLabel"])
    for proj in data["sideProjects"]:
        project_header(doc, proj["name"], proj["url"], proj["role"], proj["dates"])
        body(doc, proj["resumeBlurb"])


def render_skills(doc, data, rules):
    section_heading(doc, "Technical Skills")
    for group in data["skills"]:
        p = add_para(doc, space_after=2)
        r1 = p.add_run(f"{group['label']}: ")
        set_run(r1, size_pt=11, bold=True, color=TEXT)
        r2 = p.add_run(", ".join(group["items"]))
        set_run(r2, size_pt=11, color=MUTED)


def render_education(doc, data, rules):
    section_heading(doc, "Education")
    edu = data["education"]
    p = add_para(doc, space_after=1)
    r = p.add_run(edu["degree"])
    set_run(r, size_pt=11, bold=True, color=TEXT)
    body(doc, f"{edu['institution']}, {edu['location']}")
    meta = edu["dates"]
    if rules["includeCGPA"]:
        meta += f"  |  CGPA: {edu['cgpa']}"
    p2 = add_para(doc, space_after=2)
    r2 = p2.add_run(meta)
    set_run(r2, size_pt=10, color=ACCENT)


def render_languages(doc, data, rules):
    section_heading(doc, "Languages")
    items = [f"{l['name']} — {l['level']}" for l in data["languages"]]
    body(doc, "  ·  ".join(items))


def render_references(doc, data, rules):
    section_heading(doc, "References")
    p = add_para(doc, space_after=2)
    run = p.add_run("References available on request.")
    set_run(run, size_pt=11, italic=True, color=MUTED)


SECTION_RENDERERS = {
    "personalDetails": render_personal_details,
    "summary": render_summary,
    "experience": render_experience,
    "sideProjects": render_side_projects,
    "skills": render_skills,
    "education": render_education,
    "languages": render_languages,
    "references": render_references,
}


# ── Top-level build ────────────────────────────────────────────────────


def build_one(data, rules, out_path):
    doc = Document()

    # Page margins — A4-ish (close to the print CSS).
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.47)
        section.left_margin = Inches(0.63)
        section.right_margin = Inches(0.63)

    # Tighten the default style line spacing.
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = Pt(11)

    # Header.
    name_heading(doc, data["contact"]["name"])
    title_chip(doc, rules["titleChip"])
    contact_line(
        doc,
        [
            (data["contact"]["email"], f"mailto:{data['contact']['email']}"),
            (data["contact"]["phone"], None),
            (data["contact"]["portfolio"], f"https://{data['contact']['portfolio']}"),
        ],
    )
    contact_line(
        doc,
        [
            (rules["locationLine"], None),
            (data["contact"]["github"], f"https://{data['contact']['github']}"),
            (data["contact"]["linkedin"], f"https://{data['contact']['linkedin']}"),
        ],
    )

    # Meta strip (visa / work-mode / notice period).
    meta_items = []
    if rules.get("visaLine"):
        meta_items.append(rules["visaLine"])
    if rules.get("workModeLine"):
        meta_items.append(rules["workModeLine"])
    if rules["includeNoticePeriod"]:
        meta_items.append(f"Notice period: {data['noticePeriod']}")
    if meta_items:
        meta_line(doc, "  ·  ".join(meta_items))

    # Body sections in the order the region prescribes.
    for key in rules["sectionOrder"]:
        renderer = SECTION_RENDERERS.get(key)
        if renderer:
            renderer(doc, data, rules)

    doc.save(out_path)


def main():
    if len(sys.argv) != 3:
        print("usage: build_docx.py <json-data-file> <out-dir>", file=sys.stderr)
        sys.exit(1)

    json_path = Path(sys.argv[1])
    out_dir = Path(sys.argv[2])
    out_dir.mkdir(parents=True, exist_ok=True)

    bundle = json.loads(json_path.read_text())
    data = bundle["resume"]
    rules_by_region = bundle["RULES"]
    regions = bundle["REGIONS"]

    print(f"Building {len(regions)} region-specific DOCX resumes…")
    for region in regions:
        rules = rules_by_region[region]
        out_path = out_dir / f"anirudh-vaka-resume-{rules['fileSlug']}.docx"
        build_one(data, rules, out_path)
        size_kb = out_path.stat().st_size / 1024
        print(f"  ✓ {rules['fileSlug']:<8} → {out_path}  ({size_kb:.1f} KB)")
    print(f"Done. {len(regions)} files written to {out_dir}")


if __name__ == "__main__":
    main()
